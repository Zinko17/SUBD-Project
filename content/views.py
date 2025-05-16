from django.contrib.auth import authenticate, login, logout
from content.forms import *
from django.db import DatabaseError
from content.decorators import position_required
from django.db import transaction
from django.contrib import messages
from .models import Salaries
from django.db.models import Sum

from django.shortcuts import render, redirect, get_object_or_404
from django.db import connection
from .models import Credit, CreditPayment
from .forms import CreditForm


def home_page(request):
    return render(request, 'content/home.html')


@position_required('Технолог', 'Админ')
def raw_materials_list(request):
    rawmaterials = Rawmaterials.objects.all()
    return render(request, 'content/rawmaterials_list.html', {'rawmaterials': rawmaterials})


@position_required('Технолог', 'Админ')
def raw_material_create(request):
    if request.method == "POST":
        form = RawmaterialsForm(request.POST)
        if form.is_valid():
            rawmaterial = form.save(commit=False)
            rawmaterial.quantity = 0
            rawmaterial.totalamount = 0
            rawmaterial.save()
            form.save()
            return redirect('rawmaterials_list')
    else:
        form = RawmaterialsForm()
    return render(request, 'content/rawmaterial_form.html', {'form': form})


@position_required('Технолог', 'Админ')
def raw_material_update(request, pk):
    rawmaterial = get_object_or_404(Rawmaterials, pk=pk)
    if request.method == "POST":
        form = RawmaterialsForm(request.POST, instance=rawmaterial)
        if form.is_valid():
            form.save()
            return redirect('rawmaterials_list')
    else:
        form = RawmaterialsForm(instance=rawmaterial)
    return render(request, 'content/rawmaterial_form.html', {'form': form})


@position_required('Технолог', 'Админ')
def raw_material_delete(request, pk):
    rawmaterial = get_object_or_404(Rawmaterials, pk=pk)
    if request.method == "POST":
        rawmaterial.delete()
        return redirect('rawmaterials_list')
    return render(request, 'content/rawmaterial_confirm_delete.html', {'rawmaterial': rawmaterial})


@position_required('Технолог', 'Админ', 'Менеджер')
def finished_goods_list(request):
    finishedgoods = Finishedgoods.objects.all()
    return render(request, 'content/finishedgoods_list.html', {'finishedgoods': finishedgoods})


@position_required('Технолог', 'Админ')
def finished_goods_create(request):
    if request.method == "POST":
        form = FinishedgoodsForm(request.POST)
        if form.is_valid():
            good = form.save(commit=False)
            good.quantity = 0
            good.totalamount = 0
            good.save()
            form.save()
            return redirect('finishedgoods_list')
    else:
        form = FinishedgoodsForm()
    return render(request, 'content/finishedgoods_form.html', {'form': form})


@position_required('Технолог', 'Админ')
def finished_goods_update(request, pk):
    finishedgood = get_object_or_404(Finishedgoods, pk=pk)
    if request.method == "POST":
        form = FinishedgoodsForm(request.POST, instance=finishedgood)
        if form.is_valid():
            form.save()
            return redirect('finishedgoods_list')
    else:
        form = FinishedgoodsForm(instance=finishedgood)
    return render(request, 'content/finishedgoods_form.html', {'form': form})


@position_required('Технолог', 'Админ')
def finished_goods_delete(request, pk):
    finishedgood = get_object_or_404(Finishedgoods, pk=pk)
    if request.method == "POST":
        finishedgood.delete()
        return redirect('finishedgoods_list')
    return render(request, 'content/finishedgoods_confirm_delete.html', {'finishedgood': finishedgood})


@position_required('Технолог', 'Админ')
def unit_list(request):
    units = Units.objects.all()
    return render(request, 'content/units_list.html', {'units': units})


@position_required('Технолог', 'Админ')
def unit_create(request):
    if request.method == "POST":
        form = UnitsForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('units_list')
    else:
        form = UnitsForm()
    return render(request, 'content/unit_form.html', {'form': form})


@position_required('Технолог', 'Админ')
def unit_update(request, pk):
    unit = get_object_or_404(Units, pk=pk)
    if request.method == "POST":
        form = UnitsForm(request.POST, instance=unit)
        if form.is_valid():
            form.save()
            return redirect('units_list')
    else:
        form = UnitsForm(instance=unit)
    return render(request, 'content/unit_form.html', {'form': form})


@position_required('Технолог', 'Админ')
def unit_delete(request, pk):
    unit = get_object_or_404(Units, pk=pk)
    if request.method == "POST":
        unit.delete()
        return redirect('units _list')
    return render(request, 'content/unit_confirm_delete.html', {'unit': unit})


@position_required('Технолог', 'Админ')
def ingredient_list(request):
    ingredients = Ingredients.objects.all()
    return render(request, 'content/ingredients_list.html', {'ingredients': ingredients})


from django.core.exceptions import ValidationError


@position_required('Технолог', 'Админ')
def ingredient_create(request, product_id):
    product = get_object_or_404(Finishedgoods, pk=product_id)

    if request.method == "POST":
        form = IngredientsForm(request.POST, product_id=product_id)

        if form.is_valid():
            raw_material = form.cleaned_data['rawmaterialid']

            # Проверка на существование ингредиента
            if Ingredients.objects.filter(productid=product, rawmaterialid=raw_material).exists():
                messages.error(request, 'Этот ингредиент уже добавлен для данного продукта.')
            else:
                form.save()
                return redirect('product_ingredients', product_id=product_id)
        else:
            # Добавьте отладочный вывод ошибок формы
            messages.error(request, 'Пожалуйста, исправьте ошибки в форме. Возможно этот ингредиент уже существует')

    else:
        form = IngredientsForm(product_id=product_id)

    return render(request, 'content/ingredient_form.html', {'form': form, 'product': product})


@position_required('Технолог', 'Админ')
def ingredient_update(request, pk):
    ingredient = get_object_or_404(Ingredients, pk=pk)
    product = ingredient.productid  # Получаем связанный продукт
    product_id = product.productid

    if request.method == "POST":
        form = IngredientsForm(request.POST, instance=ingredient)
        if form.is_valid():
            form.save()
            return redirect('product_ingredients', product_id=product_id)
    else:
        form = IngredientsForm(instance=ingredient)

    return render(request, 'content/ingredient_form.html', {'form': form, 'product': product})


@position_required('Технолог', 'Админ')
def ingredient_delete(request, pk):
    ingredient = get_object_or_404(Ingredients, pk=pk)
    product = ingredient.productid  # Получаем связанный продукт
    product_id = product.productid
    if request.method == "POST":
        ingredient.delete()
        return redirect('product_ingredients', product_id=product_id)
    return render(request, 'content/ingredient_confirm_delete.html', {'ingredient': ingredient})


@position_required('Технолог', 'Админ')
def product_ingredients(request, product_id):
    product = get_object_or_404(Finishedgoods, pk=product_id)
    ingredients = Ingredients.objects.filter(productid=product)
    return render(request, 'content/product_ingredients.html', {'product': product, 'ingredients': ingredients})


@position_required('Менеджер', 'Админ')
def purchase_raw_material_view(request):
    rawmaterials = Rawmaterials.objects.all()
    employees = Employees.objects.all()
    today = date.today()

    if request.method == 'POST':
        print("POST data:", request.POST)
        form = RawMaterialPurchaseForm(request.POST)
        if form.is_valid():
            print(form.cleaned_data)
            raw_material = form.cleaned_data['rawmaterialid']
            quantity = form.cleaned_data['quantity']
            unit_price = form.cleaned_data['unit_price']
            cost = unit_price * quantity
            budget = Budget.objects.first()

            # Получаем сотрудника, который вошел в систему
            employee = Employees.objects.filter(user=request.user).first()

            if not employee:
                messages.error(request, "Не удалось найти сотрудника для текущего пользователя.")
                return redirect("purchase_raw_material")  # Название вашего пути

            try:
                with transaction.atomic():
                    # Создаем запись о закупке, привязываем сотрудника
                    purchase = form.save(commit=False)
                    purchase.totalamount = cost  # Сохраняем стоимость закупки
                    purchase.employeeid = employee  # Привязываем сотрудника
                    purchase.save()

                    # Обновляем бюджет
                    budget.totalamount -= cost
                    budget.save()

                    # Обновляем количество сырья и общую сумму
                    raw_material.quantity += quantity
                    raw_material.totalamount += cost
                    raw_material.save()

                    return redirect('rawmaterialpurchases_list')

            except ValidationError as e:
                messages.error(request, f"Ошибка: {e}")
                print(f"Ошибка: {e}")
        else:
            print("Форма не валидна!")
            for field, errors in form.errors.items():
                print(f"Ошибка в поле {field}: {errors}")

            messages.error(request, "Ошибка при оформлении закупки.")
    else:
        form = RawMaterialPurchaseForm(initial={'purchasedate': today})

    return render(request, 'content/purchase_raw_material.html', {
        'form': form,
        'rawmaterials': rawmaterials,
        'employees': employees,
        'today': today,
    })



@position_required('Менеджер', 'Админ')
def raw_material_purchases_list(request):
    purchases = Rawmaterialpurchases.objects.all()

    return render(request, 'content/purchases_list.html', {'purchases': purchases})


@position_required('Бухгалтер', 'Админ', 'Директор')
def budget_list(request):
    budgets = Budget.objects.all()
    return render(request, 'content/budget_list.html', {'budgets': budgets})


@position_required('Админ', 'Директор')
def edit_budget(request, budget_id):
    budget = get_object_or_404(Budget, pk=budget_id)
    if request.method == "POST":
        form = BudgetForm(request.POST, instance=budget)
        if form.is_valid():
            form.save()
            return redirect('budget_list')
    else:
        form = BudgetForm(instance=budget)
    return render(request, 'content/edit_budget.html', {'form': form, 'budget': budget})


@position_required('Менеджер', 'Админ')
def produce_product(request, product_id):
    product = get_object_or_404(Finishedgoods, pk=product_id)

    if request.method == "POST":
        form = ProductProductionForm(request.POST)
        if form.is_valid():
            production_quantity = form.cleaned_data.get("quantity")
            production_date = form.cleaned_data.get("productiondate")

            # Получаем сотрудника, который вошел в систему
            employee = Employees.objects.filter(user=request.user).first()

            if not employee:
                messages.error(request, "Не удалось найти сотрудника для текущего пользователя.")
                return redirect("produce_product", product_id=product_id)

            ingredients = Ingredients.objects.filter(productid=product)
            insufficient_materials = []
            total_production_cost = 0  # Итоговая сумма затрат на сырье

            for ingredient in ingredients:
                required_amount = ingredient.quantity * production_quantity
                raw_material = ingredient.rawmaterialid

                if raw_material.quantity < required_amount:
                    insufficient_materials.append(
                        f"{raw_material.name} (не хватает {required_amount - raw_material.quantity} {raw_material.unitid})"
                    )

            if insufficient_materials:
                messages.error(
                    request, "Недостаточно сырья: " + ", ".join(insufficient_materials)
                )
                return redirect("produce_product", product_id=product_id)

            try:
                with transaction.atomic():
                    for ingredient in ingredients:
                        required_amount = ingredient.quantity * production_quantity
                        raw_material = ingredient.rawmaterialid

                        # Рассчитываем стоимость использованного сырья
                        average_price = raw_material.totalamount / raw_material.quantity
                        used_cost = required_amount * average_price
                        total_production_cost += used_cost

                        # Обновляем количество и сумму сырья
                        raw_material.quantity -= required_amount
                        raw_material.totalamount -= used_cost
                        raw_material.save()

                    # Обновляем данные о продукте
                    product.quantity += production_quantity
                    product.totalamount += total_production_cost  # Добавляем стоимость производства
                    product.unit_price = product.totalamount / product.quantity  # Пересчитываем среднюю цену
                    product.save()

                    # Создаем запись о производстве и сохраняем с текущим сотрудником
                    Productproduction.objects.create(
                        productid=product,
                        quantity=production_quantity,
                        productiondate=production_date,
                        employeeid=employee  # Передаем объект сотрудника
                    )

                    return redirect("productproduction_list")
            except Exception as e:
                messages.error(request, f"Ошибка при производстве: {e}")
                return redirect("produce_product", product_id=product_id)
        else:
            messages.error(request, "Форма содержит ошибки.")
    else:
        form = ProductProductionForm()

    return render(request, "content/produce_product.html", {
        "product": product,
        "form": form,
    })


@position_required('Менеджер', 'Админ')
def productproduction_list(request):
    productproductions = Productproduction.objects.all()
    return render(request, 'content/productproduction_list.html', {'productproductions': productproductions})


import logging

logger = logging.getLogger(__name__)


@position_required('Менеджер', 'Админ')
def sell_product(request, product_id):
    finished_good = get_object_or_404(Finishedgoods, productid=product_id)

    if request.method == "POST":
        form = ProductSaleForm(request.POST, product_id=product_id)

        if form.is_valid():
            quantity = form.cleaned_data['quantity']
            sale_date = form.cleaned_data['saledate']

            # Получаем текущего сотрудника (если используется связь с пользователем)
            employee = Employees.objects.filter(user=request.user).first()

            if not employee:
                messages.error(request, "Не удалось найти сотрудника для текущего пользователя.")
                return redirect("sell_product", product_id=product_id)

            try:
                with connection.cursor() as cursor:
                    # Передаем сотрудника в хранимую процедуру
                    cursor.execute("CALL process_product_sale(%s, %s, %s, %s)",
                                   [product_id, quantity, employee.employeeid, sale_date])

                return redirect("finishedgoods_list")

            except DatabaseError as e:
                # Логируем ошибку
                logger.error(f"Ошибка при обработке продажи: {str(e)}")

                if 'Недостаточно товара на складе' in str(e):
                    messages.error(request, "Недостаточно товара на складе.")
                else:
                    messages.error(request, "Ошибка при обработке продажи. Пожалуйста, попробуйте снова.")
                return redirect("sell_product", product_id=product_id)

            except Exception as e:
                # Логируем необработанную ошибку
                logger.error(f"Необработанная ошибка: {str(e)}")
                messages.error(request, "Произошла ошибка. Пожалуйста, попробуйте позже.")
                return redirect("sell_product", product_id=product_id)

        else:
            messages.error(request, "Форма невалидна!")

    else:
        form = ProductSaleForm(product_id=product_id)

    return render(request, "content/sell_product.html", {
        "form": form,
        "product_id": product_id,
        'finished_good': finished_good,
    })



@position_required('Менеджер', 'Админ')
def productsales_list(request):
    sales = Productsales.objects.select_related('productid',
                                                'employeeid').all()  # Загружаем все продажи с данными о продукте и сотруднике
    return render(request, 'content/product_sales_list.html', {'sales': sales})


@position_required('Админ', 'Директор')
def salary_view(request):
    years = [2025, 2024, 2023, 2022, 2021, 2020]
    months = [
        (1, 'Январь'), (2, 'Февраль'), (3, 'Март'), (4, 'Апрель'),
        (5, 'Май'), (6, 'Июнь'), (7, 'Июль'), (8, 'Август'),
        (9, 'Сентябрь'), (10, 'Октябрь'), (11, 'Ноябрь'), (12, 'Декабрь')
    ]

    selected_year = None
    selected_month = None
    salaries = None
    is_calculated = False
    total_salary_payment = 0.0
    bonuses = {}

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'update_salary':
            try:
                salary_id = int(request.POST.get('salary_id'))
                new_total = float(request.POST.get('total_salary'))
                with connection.cursor() as cursor:
                    cursor.execute("CALL update_salary_amount(%s, %s)", [salary_id, new_total])
                request.session['skip_calculation'] = True
            except Exception as e:
                messages.error(request, f"Ошибка при редактировании зарплаты: {str(e)}")
            return redirect('salary_view')

        elif 'pay' in request.POST:
            selected_year = request.session.get('selected_year')
            selected_month = request.session.get('selected_month')
            if selected_year and selected_month:
                try:
                    with connection.cursor() as cursor:
                        cursor.execute("CALL pay_salaries(%s, %s)", [selected_year, selected_month])
                    messages.success(request, "Зарплата выплачена.")
                except Exception as e:
                    error_message = str(e)
                    if "недостаточно средств" in error_message.lower():
                        clean_msg = error_message.split("CONTEXT:")[0].strip()
                        messages.error(request, clean_msg)
                    else:
                        messages.error(request, "Произошла ошибка при выплате зарплаты.")
            return redirect('salary_view')

        else:
            selected_year = request.POST.get('year')
            selected_month = request.POST.get('month')

    else:
        selected_year = request.session.get('selected_year')
        selected_month = request.session.get('selected_month')

    if selected_year and selected_month:
        try:
            selected_year = int(selected_year)
            selected_month = int(selected_month)

            request.session['selected_year'] = selected_year
            request.session['selected_month'] = selected_month

            skip_calc = request.session.pop('skip_calculation', False)
            if not skip_calc:
                try:
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT calculate_salaries2(%s, %s)", [selected_year, selected_month])
                        total_salary_payment = cursor.fetchone()[0]
                except Exception as e:
                    messages.error(request, f"Ошибка при получении общей суммы: {str(e)}")
            else:
                try:
                    total_salary_payment = Salaries.objects.filter(
                        year=selected_year,
                        month=selected_month,
                        is_paid=False  # Только для невыплаченных
                    ).aggregate(Sum('total_salary'))['total_salary__sum'] or 0.0
                except Exception as e:
                    messages.error(request, f"Ошибка при ручном подсчёте суммы: {str(e)}")

            salaries = Salaries.objects.filter(year=selected_year, month=selected_month)
            is_calculated = salaries.exists()

            for salary in salaries:
                try:
                    employee_salary = salary.employeeid.salary or 0
                    bonuses[salary.salaryid] = (employee_salary * salary.bonus_percent / 100) * salary.total_activities
                except:
                    bonuses[salary.salaryid] = 0

        except ValueError:
            messages.error(request, "Неверный формат года или месяца.")

    return render(request, 'content/salary_view.html', {
        'years': years,
        'months': months,
        'salaries': salaries,
        'selected_year': selected_year,
        'selected_month': selected_month,
        'is_calculated': is_calculated,
        'total_salary': total_salary_payment,
        'bonuses': bonuses,
    })


@position_required('Бухгалтер', 'Админ', 'Директор')
def salary_list(request):
    # Получаем все записи по зарплатам
    salaries = Salaries.objects.all()

    return render(request, 'content/salary_list.html', {
        'salaries': salaries
    })


@position_required('Бухгалтер', 'Админ')
def credit_list(request):
    credits = Credit.objects.all().order_by('-credit_date')
    return render(request, 'content/credit_list.html', {
        'credits': credits,
    })


@position_required('Бухгалтер', 'Админ')
def credit_create(request):
    if request.method == 'POST':
        form = CreditForm(request.POST)
        if form.is_valid():
            cd = form.cleaned_data
            with connection.cursor() as cur:
                cur.callproc('create_credit', [
                    cd['amount'],
                    cd['interest'],
                    cd['credit_date'],
                    cd['months'],
                    cd['late_fee'],
                ])
            return redirect('credit_list')
    else:
        form = CreditForm()
    return render(request, 'content/credit_form.html', {
        'form': form,
    })


@position_required('Бухгалтер', 'Админ')
def credit_detail(request, creditid):
    credit = get_object_or_404(Credit, creditid=creditid)
    payments = CreditPayment.objects.filter(creditid=credit).order_by('payment_date')
    next_due = credit.next_due  # DateField

    if request.method == 'POST' and not credit.is_paid_off:
        form = PaymentDateForm(request.POST)
        if form.is_valid():
            pd = form.cleaned_data['payment_date']
            try:
                with connection.cursor() as cur:
                    cur.callproc('create_credit_paymentt', [creditid, pd])
            except DatabaseError as e:
                # Проверка ошибки недостаточности средств
                if 'Недостаточно средств' in str(e):
                    messages.error(request, 'Недостаточно средств в бюджете для выполнения платежа.')
                else:
                    messages.error(request, f'Произошла ошибка при обработке платежа: {str(e)}')
                return render(request, 'content/credit_detail.html', {
                    'credit': credit,
                    'payments': payments,
                    'form': form,  # передаем форму с введенными данными
                    'next_due': next_due,
                })

            return redirect('credit_detail', creditid=creditid)
        else:
            messages.error(request, 'Пожалуйста, исправьте ошибки в форме.')

    else:
        form = PaymentDateForm()

    return render(request, 'content/credit_detail.html', {
        'credit': credit,
        'payments': payments,
        'form': form,
        'next_due': next_due,
    })


def login_view(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            u = form.cleaned_data['username']
            p = form.cleaned_data['password']
            user = authenticate(request, username=u, password=p)
            if user is not None:
                login(request, user)
                return redirect('home')  # или на любую главную страницу после логина
            else:
                messages.error(request, 'Неверный логин или пароль')
    else:
        form = LoginForm()
    return render(request, 'content/login.html', {'form': form})


def logout_view(request):
    logout(request)
    return redirect('login')

@position_required('Директор', 'Админ')
def employee_list(request):
    employees = Employees.objects.all()
    return render(request, 'content/employee_list.html', {'employees': employees})

@position_required('Директор', 'Админ')
def position_list(request):
    positions = Positions.objects.all()
    return render(request, 'content/position_list.html', {'positions': positions})

import io
import os
from django.http import FileResponse
from django.shortcuts import render, redirect
from django.contrib import messages
from django.conf import settings
from django.template.loader import render_to_string
from openpyxl import Workbook
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from .forms import ReportForm
from .models import Employees
from django.db import connection
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics




PROC_MAP = {
    'purchases':  'get_raw_purchases_report',
    'production': 'get_production_report',
    'sales':      'get_sales_report',
    'salaries':   'get_salaries_report',
    'credit_pay': 'get_credit_payments_report',
}

REPORT_ROLE_MAP = {
    'purchases':  ['Менеджер', 'Админ'],
    'production': ['Технолог', 'Админ'],
    'sales':      ['Менеджер', 'Админ'],
    'salaries':   ['Админ', 'Бухгалтер', 'Директор'],
    'credit_pay': ['Бухгалтер', 'Админ'],
}

# Путь к кастомному шрифту (например, DejaVuSans)
font_path = r'content/static/content/fonts/DejaVuSans.ttf'

# Регистрация шрифта
pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))

# Маппинг столбцов для каждого типа отчета
CREDIT_PAY_COLUMNS = {
    'paymentid': 'Номер Платежа',
    'creditid': 'Номер Кредита',
    'payment_date': 'Дата Платежа',
    'total_payment': 'Общий Платеж',
    'overdue_days': 'Просроченные Дни',
    'penalty': 'Пени'
}

SALARY_COLUMNS = {
    'SalaryID': 'Номер Зарплаты',
    'EmployeeName': 'Сотрудник',
    'Year': 'Год',
    'Month': 'Месяц',
    'TotalActivities': 'Общее Количество Деятельностей',
    'BonusPercent': 'Процент Бонуса',
    'TotalSalary': 'Общая Зарплата',
    'IsPaid': 'Оплачено'
}

PURCHASES_COLUMNS = {
    'purchaseid': 'Номер закупки',
    'quantity': 'Количество',
    'totalamount': 'Общая Сумма',
    'purchasedate': 'Дата Закупки',
    'employeename': 'Сотрудник',
    'rawmaterialname': 'Сырьe'
}

SALES_COLUMNS = {
    'SaleID': 'Номер продажи',
    'Quantity': 'Количество',
    'TotalAmount': 'Общая Сумма',
    'SaleDate': 'Дата Продажи',
    'Employee': 'Сотрудник',
    'Product': 'Название продукта'
}

PRODUCTION_COLUMNS = {
    'production_id': 'Номер Производства',
    'product_name': 'Название Продукта',
    'quantity': 'Количество',
    'production_date': 'Дата Производства',
    'employee_name': 'Сотрудник'
}


# Маппинг для каждого типа отчета
COLUMN_MAP = {
    'credit_pay': CREDIT_PAY_COLUMNS,
    'salaries': SALARY_COLUMNS,
    'purchases': PURCHASES_COLUMNS,
    'sales': SALES_COLUMNS,
    'production': PRODUCTION_COLUMNS,
}

# Обработчик отчётов
@position_required('Бухгалтер','Директор', 'Админ', 'Менеджер', 'Технолог')
def report_view(request):
    form = ReportForm(request.GET or None)
    data, columns = [], []

    if form.is_valid():
        rt = form.cleaned_data['report_type']
        sd = form.cleaned_data['start_date']
        ed = form.cleaned_data['end_date']

        # Проверка профиля
        try:
            employee = request.user.employee_profile
            user_position = employee.positionid.positionname
        except Employees.DoesNotExist:
            messages.error(request, "Ваша учётная запись не связана с сотрудником.")
            return redirect('home')

        if user_position not in REPORT_ROLE_MAP.get(rt, []):
            messages.error(request, f"У вас нет права формировать отчёт «{rt}».")
            return redirect('report_view')

        # Получение данных
        proc = PROC_MAP[rt]
        with connection.cursor() as cur:
            cur.execute(f"SELECT * FROM {proc}(%s, %s)", [sd, ed])
            columns = [col[0] for col in cur.description]
            data = [dict(zip(columns, row)) for row in cur.fetchall()]

        # Заменяем столбцы на русские
        if rt in COLUMN_MAP:
            column_map = COLUMN_MAP[rt]
            columns = [column_map.get(col, col) for col in columns]  # Заменяем имена столбцов
            data = [
                {column_map.get(col, col): row[col] for col in row}
                for row in data
            ]

        fmt = request.GET.get('export_format')
        if fmt:
            if fmt == 'xlsx':
                wb = Workbook()
                ws = wb.active
                ws.append(columns)
                for row in data:
                    ws.append([row[c] for c in columns])
                stream = io.BytesIO()
                wb.save(stream)
                stream.seek(0)
                return FileResponse(stream,
                                    as_attachment=True,
                                    filename=f'report_{rt}.xlsx')




            elif fmt == 'docx':
                # Получаем директора
                director = Employees.objects.filter(positionid__positionname__icontains='Директор').first()
                # Получаем ответственного — текущего пользователя
                responsible = Employees.objects.filter(user=request.user).first()
                # Проверка на случай, если кто-то не найден
                director_name = director.fullname if director else 'Неизвестно'
                responsible_name = responsible.fullname if responsible else 'Неизвестно'
                user_position = responsible.positionid.positionname if responsible else 'Не указано'  # Позиция ответственного
                doc = Document()
                # Заголовок
                report_title = form.get_report_label()
                doc.add_heading(f'Отчёт: {report_title}', level=1)
                doc.add_paragraph(f'Период: с {sd.strftime("%d.%m.%Y")} по {ed.strftime("%d.%m.%Y")}')
                # Добавление таблицы
                table = doc.add_table(rows=1, cols=len(columns))
                hdr = table.rows[0].cells
                for i, c in enumerate(columns):
                    hdr[i].text = c
                for row in data:
                    cells = table.add_row().cells
                    for i, c in enumerate(columns):
                        cells[i].text = str(row[c])
                # Подписи
                doc.add_paragraph("\n\n")
                doc.add_paragraph(f"Директор: {director_name}, Подпись: ___________")
                doc.add_paragraph(f"{user_position}: {responsible_name}, Подпись: ___________")
                # Сохранение в поток
                stream = io.BytesIO()
                doc.save(stream)
                stream.seek(0)
                return FileResponse(stream, as_attachment=True, filename=f'report_{rt}.docx')



            elif fmt == 'pdf':
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.lib.styles import ParagraphStyle
                from reportlab.platypus import Paragraph, Spacer, Table
                from reportlab.lib.pagesizes import A4
                from reportlab.lib import colors
                from reportlab.lib.units import mm
                # Подключаем шрифт
                pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                stream = io.BytesIO()
                doc = SimpleDocTemplate(stream, pagesize=A4, rightMargin=20, leftMargin=20, topMargin=20,
                                        bottomMargin=20)
                styles = getSampleStyleSheet()
                russian_style = ParagraphStyle(
                    name='Russian',
                    fontName='DejaVuSans',
                    fontSize=8,
                    leading=10,
                    alignment=1,  # центр
                    wordWrap='CJK'
                )
                elements = []
                # Заголовок
                report_title = form.get_report_label()
                elements.append(Paragraph(f"<b>Отчёт: {report_title}</b>", russian_style))
                elements.append(
                    Paragraph(f"Период: с {sd.strftime('%d.%m.%Y')} по {ed.strftime('%d.%m.%Y')}", russian_style))
                elements.append(Spacer(1, 10))
                # Подготовка данных таблицы с оборачиванием текста
                max_table_width = A4[0] - 40  # ширина страницы минус отступы
                num_cols = len(columns)
                col_width = max_table_width / num_cols
                # Заголовки с переносом
                wrapped_headers = [Paragraph(c, russian_style) for c in columns]
                # Данные с оборачиванием текста
                wrapped_data = [
                    [Paragraph(str(row[col]), russian_style) for col in columns]
                    for row in data
                ]
                table_data = [wrapped_headers] + wrapped_data
                table = Table(table_data, repeatRows=1, colWidths=[col_width] * num_cols)
                table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans'),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#d3d3d3')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ]))
                elements.append(table)
                # Подписи
                elements.append(Spacer(1, 24))
                # Получаем данные для директора
                director = Employees.objects.filter(positionid__positionname__icontains='директор').first()
                director_name = director.fullname if director else 'Неизвестно'
                # Получаем ответственного — текущего пользователя
                responsible = Employees.objects.filter(user=request.user).first()
                responsible_name = responsible.fullname if responsible else 'Неизвестно'
                user_position = responsible.positionid.positionname if responsible else 'Не указано'
                elements.append(Paragraph(f"Директор: {director_name}, Подпись: ____________________", russian_style))
                elements.append(
                    Paragraph(f"{user_position}: {responsible_name}, Подпись: ____________________",
                              russian_style))
                doc.build(elements)
                stream.seek(0)
                return FileResponse(stream, as_attachment=True, filename=f'report_{rt}.pdf')
    return render(request, 'content/report.html', {
        'form': form,
        'data': data,
        'columns': columns,
        'request': request,
    })



